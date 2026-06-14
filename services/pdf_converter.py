"""
PDF to Word Converter v2 — High-fidelity conversion with OCR, images, tables.

Pipeline:
  PDF → PyMuPDF text extraction (blocks mode)
    ├─ Text found → layout-aware extraction with styles
    ├─ No text (scanned) → render page → OCR pipeline
    └─ Mixed → text + OCR for complete coverage
  → Image extraction + embedding
  → Table detection + Word table creation
  → python-docx generation with formatting preservation
"""

import fitz  # PyMuPDF
import docx
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
import logging
import tempfile
from PIL import Image
from io import BytesIO

from services.ocr_processor import OCRProcessor
from services.layout_analyzer import LayoutAnalyzer

logger = logging.getLogger(__name__)


class PDFConverter:
    """
    High-fidelity PDF to Word converter.

    Features:
    - Layout-aware text extraction via PyMuPDF
    - OCR for scanned pages (PaddleOCR or Tesseract)
    - Image extraction and embedding
    - Table detection and Word table creation
    - Heading/style detection
    - Progress tracking via callback
    """

    def __init__(self):
        self.ocr = OCRProcessor()
        self.layout = LayoutAnalyzer()
        logger.info(f"PDFConverter v2 initialized — OCR engine: {self.ocr.engine_name}")

    # ─── Main Conversion Entry Point ────────────────────────────

    def convert_pdf_to_word(self, pdf_path: str, output_path: str,
                            progress_callback=None) -> bool:
        """
        Convert PDF to Word document with full pipeline.

        Args:
            pdf_path: Path to input PDF
            output_path: Path for output DOCX
            progress_callback: Optional fn(percent, message) for progress

        Returns:
            bool: True if successful
        """
        try:
            self._progress(progress_callback, 5, "Opening PDF document...")
            pdf_document = fitz.open(pdf_path)
            total_pages = len(pdf_document)

            # Validate
            if total_pages == 0:
                self._progress(progress_callback, 0, "Error: PDF has no pages")
                pdf_document.close()
                return False

            if total_pages > 100:
                self._progress(progress_callback, 0, "Error: PDF exceeds 100 page limit")
                pdf_document.close()
                return False

            self._progress(progress_callback, 10, f"Analyzing {total_pages} pages...")

            # Create Word document
            doc = docx.Document()
            self._setup_document_styles(doc)

            # Set metadata
            doc.core_properties.title = os.path.basename(pdf_path)
            doc.core_properties.subject = "PDF to Word Conversion"

            # Process each page
            for page_num in range(total_pages):
                base_progress = 10 + (page_num / total_pages) * 70
                self._progress(progress_callback, base_progress,
                               f"Processing page {page_num + 1} of {total_pages}...")

                page = pdf_document.load_page(page_num)
                self._convert_page(doc, page, page_num, total_pages,
                                   progress_callback, base_progress)

            self._progress(progress_callback, 85, "Finalizing document structure...")

            # Save
            self._progress(progress_callback, 92, "Saving Word document...")
            doc.save(output_path)

            pdf_document.close()

            self._progress(progress_callback, 100, "Conversion completed successfully!")
            return True

        except Exception as e:
            logger.error(f"PDF conversion failed: {e}", exc_info=True)
            self._progress(progress_callback, 0, f"Error: {str(e)}")
            return False

    # ─── Per-Page Conversion ────────────────────────────────────

    def _convert_page(self, doc: docx.Document, page: fitz.Page,
                       page_num: int, total_pages: int,
                       progress_callback, base_progress: float):
        """Convert a single PDF page to Word content."""

        # Add page separator (except before first page)
        if page_num > 0:
            self._add_page_separator(doc, page_num + 1)

        # ── Step 1: Text extraction (block-level) ──
        text_blocks = self._extract_text_page(page)

        if text_blocks:
            # Page has extractable text — add it with formatting
            logger.debug(f"Page {page_num + 1}: {len(text_blocks)} text blocks extracted")
            self._add_text_blocks_to_doc(doc, text_blocks)
        else:
            # ── Step 2: No text → scanned page → OCR ──
            logger.debug(f"Page {page_num + 1}: No text — running OCR")
            self._progress(progress_callback, base_progress + 30,
                           f"Page {page_num + 1}: Running OCR on scanned page...")
            self._ocr_page_to_doc(doc, page, page_num)

        # ── Step 3: Extract and embed images ──
        self._extract_and_embed_images(doc, page, page_num)

        # ── Step 4: Detect and create tables ──
        self._detect_and_create_tables(doc, page, page_num)

    # ─── Text Extraction ────────────────────────────────────────

    def _extract_text_page(self, page: fitz.Page) -> list[dict]:
        """
        Extract text blocks from a page with layout information.

        Uses PyMuPDF's block-level extraction for reading order
        and falls back to simple text extraction if blocks are empty.
        """
        # Primary: block-level structured extraction
        blocks_raw = page.get_text("blocks")

        text_blocks = []
        for block in blocks_raw:
            x0, y0, x1, y1, text, block_type, block_no = block
            text = text.strip() if text else ""

            if text and block_type == 0:  # type 0 = text
                # Try to get font info from dict extraction
                font_info = self._get_block_font_info(page, (x0, y0, x1, y1))

                text_blocks.append({
                    'text': text,
                    'bbox': (x0, y0, x1, y1),
                    'block_no': block_no,
                    'font_size': font_info.get('size', 0),
                    'font_name': font_info.get('font', ''),
                    'is_bold': font_info.get('bold', False),
                    'is_italic': font_info.get('italic', False),
                })

        # If blocks gave nothing, try dict mode
        if not text_blocks:
            text_blocks = self._extract_text_dict_mode(page)

        return text_blocks

    def _extract_text_dict_mode(self, page: fitz.Page) -> list[dict]:
        """Fallback text extraction using dict mode for better structure."""
        blocks = []
        try:
            text_dict = page.get_text("dict")
            for block in text_dict.get("blocks", []):
                if block.get("type") != 0:  # not text
                    continue

                block_text_parts = []
                font_info = {}
                block_bbox = list(block["bbox"])

                for line in block.get("lines", []):
                    line_text = []
                    for span in line.get("spans", []):
                        line_text.append(span.get("text", ""))
                        if not font_info:
                            font_info = {
                                'size': span.get('size', 0),
                                'font': span.get('font', ''),
                                'bold': bool(span.get('flags', 0) & 2**3),
                                'italic': bool(span.get('flags', 0) & 2**1),
                            }
                    block_text_parts.append(''.join(line_text))

                full_text = '\n'.join(block_text_parts).strip()
                if full_text:
                    blocks.append({
                        'text': full_text,
                        'bbox': tuple(block_bbox),
                        'block_no': block.get("number", 0),
                        **font_info,
                    })
        except Exception as e:
            logger.debug(f"Dict mode extraction failed: {e}")

        return blocks

    def _get_block_font_info(self, page: fitz.Page, bbox: tuple) -> dict:
        """Get font information for a region on the page."""
        try:
            text_dict = page.get_text("dict", clip=bbox)
            for block in text_dict.get("blocks", []):
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        if span.get("text", "").strip():
                            return {
                                'size': span.get('size', 0),
                                'font': span.get('font', ''),
                                'bold': bool(span.get('flags', 0) & 2**3),
                                'italic': bool(span.get('flags', 0) & 2**1),
                            }
        except Exception:
            pass
        return {}

    # ─── OCR Pipeline ───────────────────────────────────────────

    def _ocr_page_to_doc(self, doc: docx.Document, page: fitz.Page, page_num: int):
        """
        Render a page to image and run OCR to extract text.
        Inserts the recognized text into the Word document.
        """
        try:
            # Render page to image at high resolution
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Try block-level OCR (preserves layout better)
            ocr_blocks = self.ocr.extract_text_blocks(img)

            if ocr_blocks:
                # Sort blocks top-to-bottom, left-to-right for reading order
                ocr_blocks.sort(key=lambda b: (b['center_y'], b['center_x']))

                # Group into paragraphs by vertical gaps
                paragraphs = self._group_ocr_blocks_into_paragraphs(ocr_blocks)

                for para_text in paragraphs:
                    if para_text.strip():
                        paragraph = doc.add_paragraph(para_text.strip())
                        # OCR text gets slightly smaller font since it's reconstructed
                        for run in paragraph.runs:
                            run.font.size = Pt(11)

                logger.debug(f"OCR extracted {len(paragraphs)} paragraphs from page {page_num + 1}")
            else:
                # Fallback to full-page OCR
                text, confidence = self.ocr.extract_text(img)
                if text.strip():
                    doc.add_paragraph(text.strip())
                    logger.debug(f"OCR full-page extraction: {confidence:.1f}% confidence")

                    # Add confidence note for very low-quality results
                    if confidence < 70:
                        note = doc.add_paragraph()
                        note.add_run(
                            f"[Note: Low OCR confidence ({confidence:.1f}%). "
                            "Please review extracted text carefully.]"
                        ).italic = True

        except Exception as e:
            logger.error(f"OCR page processing failed: {e}")
            # Last resort: add a warning
            doc.add_paragraph(
                f"[Scanned content on page — could not extract text. "
                f"Please use a text-based PDF for best results.]"
            )

    def _group_ocr_blocks_into_paragraphs(self, blocks: list[dict]) -> list[str]:
        """
        Group OCR text blocks into logical paragraphs based on vertical spacing.

        Blocks that are close together vertically get merged into paragraphs.
        Large vertical gaps indicate paragraph breaks.
        """
        if not blocks:
            return []

        paragraphs = []
        current_para = blocks[0]['text']
        prev_y = blocks[0]['center_y']

        # Estimate line height from average block height
        heights = [b['bbox'][3] - b['bbox'][1] for b in blocks]
        avg_line_height = sum(heights) / len(heights) if heights else 15

        for block in blocks[1:]:
            gap = block['center_y'] - prev_y

            if gap > avg_line_height * 2.5:
                # Large gap → new paragraph
                paragraphs.append(current_para)
                current_para = block['text']
            elif gap > avg_line_height * 1.2:
                # Medium gap → new line within paragraph
                current_para += '\n' + block['text']
            else:
                # Small gap → same line
                current_para += ' ' + block['text']

            prev_y = block['center_y']

        # Don't forget the last one
        paragraphs.append(current_para)
        return paragraphs

    # ─── Image Extraction ───────────────────────────────────────

    def _extract_and_embed_images(self, doc: docx.Document, page: fitz.Page,
                                   page_num: int):
        """
        Extract images from the page and embed them in the Word document.
        """
        try:
            image_list = page.get_images(full=True)

            for img_idx, img_info in enumerate(image_list):
                xref = img_info[0]  # Image xref number
                base_image = page.parent.extract_image(xref)

                if base_image:
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]  # e.g., 'jpeg', 'png'

                    try:
                        # Open and optimize image
                        pil_img = Image.open(BytesIO(image_bytes))

                        # Limit max size to avoid huge DOCX files
                        max_dim = 1200
                        if pil_img.width > max_dim or pil_img.height > max_dim:
                            ratio = max_dim / max(pil_img.width, pil_img.height)
                            new_size = (int(pil_img.width * ratio), int(pil_img.height * ratio))
                            pil_img = pil_img.resize(new_size, Image.LANCZOS)

                        # Save optimized image to bytes
                        img_buffer = BytesIO()
                        pil_img.save(img_buffer, format='PNG', optimize=True)
                        img_buffer.seek(0)

                        # Add to document
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        run = paragraph.add_run()
                        run.add_picture(
                            img_buffer,
                            width=Inches(min(6.0, pil_img.width / 150)),
                        )

                        # Add caption
                        caption = doc.add_paragraph()
                        caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        caption_run = caption.add_run(
                            f"[Image from page {page_num + 1}]"
                        )
                        caption_run.font.size = Pt(9)
                        caption_run.font.color.rgb = RGBColor(128, 128, 128)
                        caption_run.italic = True

                        logger.debug(f"Extracted image {img_idx + 1} from page {page_num + 1}")

                    except Exception as e:
                        logger.error(f"Failed to process image: {e}")

        except Exception as e:
            logger.debug(f"Image extraction skipped for page {page_num + 1}: {e}")

    # ─── Table Detection & Creation ─────────────────────────────

    def _detect_and_create_tables(self, doc: docx.Document, page: fitz.Page,
                                   page_num: int):
        """
        Detect tables on the page and create real Word tables.
        """
        try:
            # Use PyMuPDF's built-in table detection (v1.23+)
            tabs = page.find_tables(
                strategy='lines',
                vertical_strategy='lines',
                horizontal_strategy='lines',
            )

            if tabs and tabs.tables:
                for table in tabs.tables:
                    cells = table.extract()

                    if not cells:
                        continue

                    # Add spacing before table
                    doc.add_paragraph()

                    # Create Word table
                    num_rows = len(cells)
                    num_cols = max(len(row) for row in cells) if cells else 1

                    word_table = doc.add_table(rows=num_rows, cols=num_cols, style='Table Grid')
                    word_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    # Fill cells
                    for row_idx, row_cells in enumerate(cells):
                        for col_idx, cell_text in enumerate(row_cells):
                            if col_idx < num_cols:
                                cell = word_table.cell(row_idx, col_idx)
                                cell.text = str(cell_text).strip() if cell_text else ''

                                # Style header row (first row)
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
                                            run.font.size = Pt(10)

                    # Add spacing after table
                    doc.add_paragraph()
                    logger.debug(f"Created table ({num_rows}x{num_cols}) from page {page_num + 1}")

        except Exception as e:
            # Table detection requires PyMuPDF >= 1.23
            logger.debug(f"Table detection skipped: {e}")

    # ─── Text Blocks → Word Document ────────────────────────────

    def _add_text_blocks_to_doc(self, doc: docx.Document, blocks: list[dict]):
        """
        Add extracted text blocks to the Word document with formatting.

        Applies heading detection and basic style preservation.
        """
        # Detect headings
        self._mark_headings(blocks)

        for block in blocks:
            text = block['text'].strip()
            if not text or len(text) < 2:
                continue

            # Split into sub-paragraphs if the block has multiple paragraphs
            sub_paras = self._split_block_into_paragraphs(text)
            first = True

            for sub_text in sub_paras:
                sub_text = sub_text.strip()
                if not sub_text or len(sub_text) < 2:
                    continue

                paragraph = doc.add_paragraph()

                # Determine style
                is_heading = block.get('is_heading', False)
                heading_level = block.get('heading_level', 0)

                if is_heading:
                    if first:
                        # Apply heading style
                        if heading_level == 1:
                            paragraph.style = doc.styles['Heading 1']
                        elif heading_level == 2:
                            paragraph.style = doc.styles['Heading 2']
                        else:
                            paragraph.style = doc.styles['Heading 3']
                    else:
                        # Continuation of heading — use normal style
                        pass

                # Add text with formatting
                run = paragraph.add_run(sub_text)

                if block.get('font_size', 0) > 0:
                    run.font.size = Pt(block['font_size'])
                if block.get('is_bold', False):
                    run.bold = True
                if block.get('is_italic', False):
                    run.italic = True

                first = False

    def _mark_headings(self, blocks: list[dict]):
        """
        Mark blocks that are likely headings based on heuristics.
        Modifies blocks in-place with 'is_heading' and 'heading_level' keys.
        """
        if not blocks:
            return

        # Gather font sizes
        font_sizes = [b.get('font_size', 0) for b in blocks if b.get('font_size', 0) > 0]
        if not font_sizes:
            return

        avg_size = sum(font_sizes) / len(font_sizes)
        max_size = max(font_sizes)

        for block in blocks:
            text = block['text'].strip()
            text_len = len(text)
            font_size = block.get('font_size', 0)
            is_bold = block.get('is_bold', False)

            if not font_size:
                continue

            # H1-level: significantly larger than average
            if font_size >= max_size * 0.9 and text_len < 150:
                block['is_heading'] = True
                block['heading_level'] = 1
            # H2-level: bold + larger than average
            elif is_bold and font_size > avg_size * 1.1 and text_len < 120:
                block['is_heading'] = True
                block['heading_level'] = 2
            # H3-level: bold + short text
            elif is_bold and text_len < 80 and not text.endswith('.'):
                block['is_heading'] = True
                block['heading_level'] = 3
            # All-caps short text at or above average size
            elif text_len < 60 and text.isupper() and font_size >= avg_size:
                block['is_heading'] = True
                block['heading_level'] = 3
            else:
                block['is_heading'] = False
                block['heading_level'] = 0

    def _split_block_into_paragraphs(self, text: str) -> list[str]:
        """
        Split a text block into logical paragraphs.

        Respects existing paragraph breaks (double newlines) and
        avoids creating too many tiny paragraphs.
        """
        # Split on double newlines (true paragraph breaks)
        raw_paras = text.split('\n\n')

        # Clean each paragraph
        cleaned = []
        for para in raw_paras:
            para = para.strip()
            # Replace single newlines with spaces within a paragraph
            para = para.replace('\n', ' ')
            # Normalize whitespace
            para = ' '.join(para.split())
            if para and len(para) > 2:
                cleaned.append(para)

        # Merge very short paragraphs with adjacent ones
        merged = []
        buffer = ""

        for para in cleaned:
            if len(para) < 40 and buffer and not para.endswith('.') and not para.endswith(':'):
                buffer += ' ' + para
            else:
                if buffer:
                    merged.append(buffer)
                buffer = para

        if buffer:
            merged.append(buffer)

        return merged if merged else [text]

    # ─── Page Separator ─────────────────────────────────────────

    def _add_page_separator(self, doc: docx.Document, page_number: int):
        """Add a subtle page separator between pages."""
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(f"─ Page {page_number} ─")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(180, 180, 180)
        run.italic = True

    # ─── Document Setup ─────────────────────────────────────────

    def _setup_document_styles(self, doc: docx.Document):
        """Configure Word document styles for better output."""
        # Default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Heading styles
        for i, (size, color) in enumerate([
            (16, RGBColor(0x1a, 0x1a, 0x1a)),  # H1
            (14, RGBColor(0x2d, 0x2d, 0x2d)),  # H2
            (12, RGBColor(0x40, 0x40, 0x40)),  # H3
        ], start=1):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.size = Pt(size)
            heading_style.font.color.rgb = color
            heading_style.font.bold = True

        # Paragraph spacing
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.space_before = Pt(2)

    # ─── Utilities ──────────────────────────────────────────────

    @staticmethod
    def _progress(callback, percent: float, message: str):
        """Call progress callback if set."""
        if callback:
            try:
                callback(int(percent), message)
            except Exception:
                pass

    def validate_pdf(self, pdf_path: str) -> tuple[bool, str]:
        """
        Validate a PDF file before conversion.

        Returns:
            (is_valid, message)
        """
        try:
            doc = fitz.open(pdf_path)
            page_count = len(doc)

            if page_count == 0:
                doc.close()
                return False, "PDF file contains no pages"

            if page_count > 100:
                doc.close()
                return False, "PDF file has too many pages (maximum 100)"

            # Quick check for encryption
            if doc.is_encrypted:
                doc.close()
                return False, "PDF file is password-protected — cannot convert"

            doc.close()
            return True, f"Valid PDF with {page_count} pages"

        except Exception as e:
            return False, f"Invalid PDF file: {str(e)}"
